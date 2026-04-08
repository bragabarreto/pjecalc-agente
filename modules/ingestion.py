# modules/ingestion.py — Módulo de Ingestão e Leitura da Sentença
# Manual Técnico PJE-Calc, Seção 2

from __future__ import annotations

import io
import re
import unicodedata
from pathlib import Path
from typing import Any

import chardet


# ── Funções públicas ──────────────────────────────────────────────────────────

def ler_documento(caminho: str | Path) -> dict[str, Any]:
    """
    Ponto de entrada principal do módulo.
    Detecta o formato e extrai o texto normalizado + metadados.

    Retorna:
        {
            "texto": str,
            "paginas": list[str],        # texto por página (PDFs)
            "formato": str,              # "pdf_nativo" | "pdf_ocr" | "docx" | "txt"
            "encoding_original": str,
            "num_paginas": int,
            "alertas": list[str],        # avisos para o usuário
        }
    """
    caminho = Path(caminho)
    sufixo = caminho.suffix.lower()
    alertas: list[str] = []

    if sufixo == ".pdf":
        resultado = _ler_pdf(caminho, alertas)
    elif sufixo in (".docx", ".doc"):
        resultado = _ler_docx(caminho, alertas)
    elif sufixo == ".txt":
        resultado = _ler_txt(caminho, alertas)
    else:
        raise ValueError(f"Formato não suportado: {sufixo}")

    resultado["texto"] = normalizar_texto(resultado["texto"])
    resultado["alertas"] = alertas
    return resultado


# ── Leitura por formato ───────────────────────────────────────────────────────

_PDF_MAX_PAGINAS = 40  # limite por documento para não estourar memória na nuvem

def _ler_pdf(caminho: Path, alertas: list[str]) -> dict[str, Any]:
    """Tenta extração nativa; recorre ao OCR se o texto for insuficiente."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("Instale pdfplumber: pip install pdfplumber")

    paginas: list[str] = []
    with pdfplumber.open(str(caminho)) as pdf:
        total = len(pdf.pages)
        if total > _PDF_MAX_PAGINAS:
            alertas.append(
                f"PDF com {total} páginas — processando apenas as primeiras {_PDF_MAX_PAGINAS}."
            )
        for pagina in pdf.pages[:_PDF_MAX_PAGINAS]:
            texto_pg = pagina.extract_text() or ""
            paginas.append(texto_pg)

    texto_total = "\n\n".join(paginas)

    # Heurística: se texto extraído < 100 chars por página → provavelmente OCR
    media_chars = len(texto_total) / max(len(paginas), 1)
    if media_chars < 100:
        alertas.append(
            "PDF parece ser escaneado. Aplicando OCR — verifique o resultado."
        )
        return _ler_pdf_ocr(caminho, alertas)

    return {
        "texto": texto_total,
        "paginas": paginas,
        "formato": "pdf_nativo",
        "encoding_original": "utf-8",
        "num_paginas": len(paginas),
    }


def _ler_pdf_ocr(caminho: Path, alertas: list[str]) -> dict[str, Any]:
    """Converte páginas do PDF para imagem e aplica OCR com tesseract."""
    try:
        import pytesseract
        from pdf2image import convert_from_path
        from PIL import Image
    except ImportError:
        raise ImportError(
            "Para PDFs escaneados, instale: pip install pytesseract pdf2image pillow"
        )

    from config import OCR_DPI, OCR_CONFIDENCE_MIN, OCR_LANG

    imagens = convert_from_path(str(caminho), dpi=OCR_DPI)
    paginas: list[str] = []

    for i, img in enumerate(imagens, start=1):
        dados = pytesseract.image_to_data(
            img, lang=OCR_LANG, output_type=pytesseract.Output.DICT
        )
        confs = [int(c) for c in dados["conf"] if c != "-1"]
        media_conf = sum(confs) / max(len(confs), 1)

        if media_conf < OCR_CONFIDENCE_MIN:
            alertas.append(
                f"Página {i}: confiança OCR baixa ({media_conf:.0f}%). "
                "Revise o texto extraído antes de prosseguir."
            )

        texto_pg = pytesseract.image_to_string(img, lang=OCR_LANG)
        paginas.append(texto_pg)

    # Calcular confiança média geral do OCR
    _conf_geral = sum(confs) / max(len(confs), 1) if confs else 0

    resultado = {
        "texto": "\n\n".join(paginas),
        "paginas": paginas,
        "formato": "pdf_ocr",
        "encoding_original": "ocr",
        "num_paginas": len(imagens),
    }

    # Propagar flag de baixa confiança para downstream (extração pode reduzir confiança)
    if _conf_geral < OCR_CONFIDENCE_MIN:
        resultado["_ocr_baixa_confianca"] = True
        resultado["_ocr_confianca_media"] = round(_conf_geral, 1)

    return resultado


def _ler_docx(caminho: Path, alertas: list[str]) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Instale python-docx: pip install python-docx")

    doc = Document(str(caminho))
    blocos: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            blocos.append(para.text)

    # Incluir texto de tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
            if celulas:
                blocos.append(" | ".join(celulas))

    texto = "\n".join(blocos)
    return {
        "texto": texto,
        "paginas": [texto],
        "formato": "docx",
        "encoding_original": "utf-8",
        "num_paginas": 1,
    }


def _ler_txt(caminho: Path, alertas: list[str]) -> dict[str, Any]:
    raw = caminho.read_bytes()
    det = chardet.detect(raw)
    enc = det.get("encoding") or "utf-8"
    texto = raw.decode(enc, errors="replace")
    return {
        "texto": texto,
        "paginas": [texto],
        "formato": "txt",
        "encoding_original": enc,
        "num_paginas": 1,
    }


# ── Normalização ──────────────────────────────────────────────────────────────

def normalizar_texto(texto: str) -> str:
    """
    Pipeline de normalização conforme Manual Técnico, Seção 2.2:
    1. Converter para UTF-8 (já feito na leitura)
    2. Remover cabeçalhos/rodapés repetitivos de tribunal
    3. Normalizar datas para DD/MM/AAAA
    4. Normalizar valores monetários
    5. Remover hifenação de quebra de linha
    6. Segmentar em blocos lógicos
    """
    # Unicode NFC
    texto = unicodedata.normalize("NFC", texto)

    # Remover hifenação de quebra de linha: "pala-\nvra" → "palavra"
    texto = re.sub(r"(\w)-\n(\w)", r"\1\2", texto)

    # Normalizar quebras de linha múltiplas
    texto = re.sub(r"\n{3,}", "\n\n", texto)

    # Remover cabeçalhos/rodapés típicos de tribunal (linhas repetitivas)
    texto = _remover_cabecalhos_repetitivos(texto)

    # Normalizar espaços múltiplos
    texto = re.sub(r"[ \t]{2,}", " ", texto)

    return texto.strip()


def _remover_cabecalhos_repetitivos(texto: str) -> str:
    """Remove linhas que aparecem mais de 3 vezes (tipicamente cabeçalhos de página)."""
    linhas = texto.split("\n")
    contagem: dict[str, int] = {}
    for linha in linhas:
        chave = linha.strip()
        if chave:
            contagem[chave] = contagem.get(chave, 0) + 1

    filtradas = []
    for linha in linhas:
        chave = linha.strip()
        # Manter se aparece ≤ 3 vezes OU se é uma linha significativa (>30 chars)
        if not chave or contagem[chave] <= 3 or len(chave) > 30:
            filtradas.append(linha)

    return "\n".join(filtradas)


def normalizar_valor(texto: str) -> float:
    """
    Converte string monetária brasileira para float.
    Exemplo: 'R$ 12.345,67' → 12345.67
    """
    limpo = re.sub(r"R\$\s*", "", texto)
    limpo = limpo.replace(".", "").replace(",", ".")
    return float(limpo.strip())


def normalizar_data(texto: str) -> str | None:
    """
    Tenta normalizar diferentes formatos de data para DD/MM/AAAA.
    Suporta: '01/03/2018', '01-03-2018', '01.03.2018',
             'primeiro de março de 2018', '1º de março de 2018'.
    Retorna None se não conseguir reconhecer.
    """
    # Formato numérico direto
    m = re.match(r"(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})", texto.strip())
    if m:
        d, mes, a = m.groups()
        return f"{int(d):02d}/{int(mes):02d}/{a}"

    # Formato por extenso
    MESES = {
        "janeiro": "01", "fevereiro": "02", "março": "03",
        "abril": "04", "maio": "05", "junho": "06",
        "julho": "07", "agosto": "08", "setembro": "09",
        "outubro": "10", "novembro": "11", "dezembro": "12",
    }
    m = re.search(
        r"(\d{1,2})(?:º|°)?\s+de\s+(\w+)\s+de\s+(\d{4})",
        texto.lower(),
    )
    if m:
        d, nome_mes, a = m.groups()
        mes = MESES.get(nome_mes)
        if mes:
            return f"{int(d):02d}/{mes}/{a}"

    return None


def segmentar_sentenca(texto: str) -> dict[str, str]:
    """
    Segmenta o texto da sentença em blocos lógicos:
    relatório, fundamentação e dispositivo.
    Conforme Manual, Seção 2.3.
    """
    padroes = {
        "relatorio": r"(?i)(relat[oó]rio|vistos|trata-se|cuida-se)",
        "fundamentacao": r"(?i)(fundamenta[çc][aã]o|méritos?|análise|aprecio)",
        "dispositivo": r"(?i)(pelo\s+exposto|ante\s+o\s+exposto|dispositivo|decis[aã]o|diante\s+do\s+exposto|condeno|defiro|julgo\s+procedente)",
    }

    posicoes: dict[str, int] = {}
    for bloco, padrao in padroes.items():
        m = re.search(padrao, texto)
        if m:
            posicoes[bloco] = m.start()

    # Ordenar blocos por posição
    partes_ordenadas = sorted(posicoes.items(), key=lambda x: x[1])
    resultado: dict[str, str] = {}

    for i, (nome, inicio) in enumerate(partes_ordenadas):
        fim = partes_ordenadas[i + 1][1] if i + 1 < len(partes_ordenadas) else len(texto)
        resultado[nome] = texto[inicio:fim].strip()

    # Garantir que o dispositivo seja o último (sentença com múltiplos dispositivos)
    if "dispositivo" in resultado:
        # Usar o ÚLTIMO dispositivo encontrado (acórdãos)
        matches = list(re.finditer(padroes["dispositivo"], texto))
        if len(matches) > 1:
            ultimo = matches[-1].start()
            resultado["dispositivo"] = texto[ultimo:].strip()

    return resultado
